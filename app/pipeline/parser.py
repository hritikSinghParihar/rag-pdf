import os
import fitz
import pymupdf4llm
from bs4 import BeautifulSoup
from docx import Document as DocxDocument
# import easyocr (lazy loaded)
import numpy as np
from PIL import Image
from typing import List, Dict, Any


class FileParser:
    def __init__(self):
        self._reader = None

    @property
    def reader(self):
        if self._reader is None:
            import easyocr
            self._reader = easyocr.Reader(['en'])
        return self._reader

    def parse(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Detect file type and parse accordingly.
        Returns a list of dicts: [{"text": str, "page": int}]
        """
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext == '.pdf':
            return self.parse_pdf(file_path)
        elif ext in ['.html', '.htm']:
            return self.parse_html(file_path)
        elif ext == '.docx':
            return self.parse_docx(file_path)
        elif ext == '.txt':
            return self.parse_txt(file_path)
        elif ext in ['.png', '.jpg', '.jpeg', '.tiff', '.bmp', '.webp']:
            return self.parse_image(file_path)
        else:
            raise ValueError(f"Unsupported file extension: {ext}")

    def parse_pdf(self, file_path: str) -> List[Dict[str, Any]]:
        """Extract markdown from PDF using pymupdf4llm."""
        md_pages = pymupdf4llm.to_markdown(file_path, page_chunks=True)
        return [{"text": p["text"], "page": i + 1} for i, p in enumerate(md_pages)]

    def parse_html(self, file_path: str) -> List[Dict[str, Any]]:
        """Extract clean text from HTML using BeautifulSoup."""
        with open(file_path, 'r', encoding='utf-8') as f:
            html_content = f.read()
        
        soup = BeautifulSoup(html_content, 'lxml')
        
        # Remove noisy elements
        for element in soup(['script', 'style', 'header', 'footer', 'nav', 'aside', 'form']):
            element.decompose()
            
        text = soup.get_text(separator='\n')
        # Clean up excessive whitespace
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        clean_text = '\n'.join(lines)
        
        return [{"text": clean_text, "page": 1}]

    def parse_docx(self, file_path: str) -> List[Dict[str, Any]]:
        """Extract text from DOCX."""
        doc = DocxDocument(file_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        
        # Also extract table content
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                full_text.append(" | ".join(row_text))
                
        return [{"text": "\n".join(full_text), "page": 1}]

    def parse_txt(self, file_path: str) -> List[Dict[str, Any]]:
        """Extract text from TXT."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return [{"text": f.read(), "page": 1}]
        except UnicodeDecodeError:
            with open(file_path, 'r', encoding='latin-1') as f:
                return [{"text": f.read(), "page": 1}]

    def parse_image(self, file_path: str) -> List[Dict[str, Any]]:
        """Extract text from Image using EasyOCR."""
        results = self.reader.readtext(file_path)
        # results is a list of [box, text, confidence]
        text = " ".join([res[1] for res in results])
        return [{"text": text, "page": 1}]

file_parser = FileParser()
