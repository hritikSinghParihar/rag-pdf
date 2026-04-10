import os
import logging
from pathlib import Path
from typing import List, Dict, Any, Optional
from concurrent.futures import ThreadPoolExecutor

import pymupdf4llm
from bs4 import BeautifulSoup
from docx import Document as DocxDocument

from config import config
from ocr_modules import VisionOCR
from structure_builder import StructureBuilder
from utils import is_pdf_scanned, pdf_to_images, preprocess_image

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".html", ".htm", ".docx", ".png", ".jpg", ".jpeg"}

# ─────────────────────────────────────────────
# OCR & Vision Processors
# ─────────────────────────────────────────────

def process_single_image(image_path: str, source_name: str, page_num: int = 1) -> List[Dict[str, Any]]:
    """Process a single image through Vision OCR and return chunks."""
    try:
        logger.info(f"Performing OCR on {image_path} (Source: {source_name})")
        ocr = VisionOCR()
        
        # Optional: preprocess
        # processed_path = preprocess_image(image_path)
        
        ocr_data = ocr.extract_structured_data(image_path)
        if "error" in ocr_data:
            logger.error(f"OCR error on {image_path}: {ocr_data['error']}")
            return []
            
        # Reconstruct structure
        doc_structure = StructureBuilder.build_from_ocr(ocr_data)
        doc_structure.metadata["page_number"] = page_num
        doc_structure.metadata["source"] = source_name
        
        chunks = doc_structure.get_chunks()
        
        # Format for vector store (ensure 'text' is in metadata for some systems, 
        # or just return as is if the caller handles it)
        formatted_chunks = []
        for chunk in chunks:
            data = chunk["metadata"].copy()
            data["text"] = chunk["text"]
            data["chunk_id"] = f"{source_name}_p{page_num}_c{len(formatted_chunks)}"
            formatted_chunks.append(data)
            
        return formatted_chunks
    except Exception as e:
        logger.error(f"Failed to OCR {image_path}: {e}")
        return []

# ─────────────────────────────────────────────
# Per-format loaders
# ─────────────────────────────────────────────

def load_pdf(path: str) -> List[Dict[str, Any]]:
    """
    Extract text as Markdown using PyMuPDF4LLM. 
    Falls back to OCR if the PDF appears to be scanned.
    """
    if is_pdf_scanned(path):
        logger.info(f"PDF {path} appears scanned. Falling back to OCR.")
        image_paths = pdf_to_images(path)
        all_chunks = []
        
        # Parallel OCR for large scanned PDFs
        with ThreadPoolExecutor(max_workers=config.max_ocr_workers if hasattr(config, 'max_ocr_workers') else 4) as executor:
            futures = [
                executor.submit(process_single_image, img_p, os.path.basename(path), i+1) 
                for i, img_p in enumerate(image_paths)
            ]
            for future in futures:
                all_chunks.extend(future.result())
        
        # Cleanup temp images
        for img_p in image_paths:
            try:
                os.remove(img_p)
                # Also remove processed if exists
                if os.path.exists(img_p.replace(".png", "_processed.png")):
                    os.remove(img_p.replace(".png", "_processed.png"))
            except:
                pass
        
        return all_chunks

    # Standard PDF text extraction
    md_pages = pymupdf4llm.to_markdown(path, page_chunks=True)
    return [
        {
            "text": page_data["text"], 
            "page": i + 1,
            "source": os.path.basename(path),
            "chunk_id": f"{os.path.basename(path)}_p{i+1}_c0"
        }
        for i, page_data in enumerate(md_pages)
    ]


def load_image(path: str) -> List[Dict[str, Any]]:
    """Process an image file through OCR."""
    return process_single_image(path, os.path.basename(path))


def load_txt(path: str) -> List[Dict[str, Any]]:
    """Read plain-text file and split into virtual pages (~3 000 chars each)."""
    text = Path(path).read_text(encoding="utf-8", errors="replace")
    page_size = 3000
    pages = []
    base_name = os.path.basename(path)
    for i, start in enumerate(range(0, max(len(text), 1), page_size)):
        chunk = text[start : start + page_size].strip()
        if chunk:
            pages.append({
                "text": chunk, 
                "page": i + 1, 
                "source": base_name,
                "chunk_id": f"{base_name}_p{i+1}_c0"
            })
    return pages


def load_html(path: str) -> List[Dict[str, Any]]:
    """Parse HTML with BeautifulSoup, strip boilerplate, return virtual pages."""
    html = Path(path).read_text(encoding="utf-8", errors="replace")
    soup = BeautifulSoup(html, "lxml")
    for tag in soup(["script", "style", "nav", "footer", "header", "aside", "form"]):
        tag.decompose()
    text = soup.get_text(separator="\n", strip=True)
    lines = [ln for ln in text.splitlines() if ln.strip()]
    text = "\n".join(lines)
    page_size = 3000
    pages = []
    base_name = os.path.basename(path)
    for i, start in enumerate(range(0, max(len(text), 1), page_size)):
        chunk = text[start : start + page_size].strip()
        if chunk:
            pages.append({
                "text": chunk, 
                "page": i + 1, 
                "source": base_name,
                "chunk_id": f"{base_name}_p{i+1}_c0"
            })
    return pages


def load_docx(path: str) -> List[Dict[str, Any]]:
    doc = DocxDocument(path)
    blocks: List[str] = []
    for para in doc.paragraphs:
        txt = para.text.strip()
        if txt:
            blocks.append(txt)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                blocks.append(row_text)
    page_size = 20
    pages = []
    base_name = os.path.basename(path)
    for i in range(0, max(len(blocks), 1), page_size):
        chunk = "\n".join(blocks[i : i + page_size]).strip()
        if chunk:
            pages.append({
                "text": chunk, 
                "page": i // page_size + 1, 
                "source": base_name,
                "chunk_id": f"{base_name}_p{i // page_size + 1}_c0"
            })
    return pages


# ─────────────────────────────────────────────
# Unified entry point
# ─────────────────────────────────────────────

_LOADERS = {
    ".pdf": load_pdf,
    ".txt": load_txt,
    ".html": load_html,
    ".htm": load_html,
    ".docx": load_docx,
    ".png": load_image,
    ".jpg": load_image,
    ".jpeg": load_image,
}


def ingest_files(file_paths: List[str]) -> List[Dict[str, Any]]:
    all_pages: List[Dict[str, Any]] = []
    for path in file_paths:
        if not os.path.exists(path):
            continue
        ext = Path(path).suffix.lower()
        loader = _LOADERS.get(ext)
        if loader:
            try:
                pages = loader(path)
                all_pages.extend(pages)
            except Exception as e:
                logger.error(f"Failed to ingest {path}: {e}")
    return all_pages

def ingest_pdfs(paths: List[str]) -> List[Dict[str, Any]]:
    return ingest_files(paths)
