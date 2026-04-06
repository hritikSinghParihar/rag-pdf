from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_complex_docx(path):
    doc = Document()
    
    # Title
    title = doc.add_heading('WisiPay Technical Specification: Smart Routing Engine v3.0', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Section 1: Abstract
    doc.add_heading('1. Abstract', level=1)
    doc.add_paragraph(
        "This document outlines the architectural design and implementation details "
        "of the Smart Routing Engine (SRE) v3.0. The SRE is responsible for dynamically "
        "selecting the optimal payment aggregator (PA) for every transaction based on "
        "cost, success rate, and provider latency."
    )
    
    # Section 2: Architecture Components
    doc.add_heading('2. System Components', level=1)
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Component Name'
    hdr_cells[1].text = 'Internal ID'
    hdr_cells[2].text = 'Function'
    
    components = [
        ('Ingress Gateway', 'GW-101', 'Handles incoming merchant API requests via REST/gRPC.'),
        ('Health Monitor', 'HM-202', 'Continuously pings Bank/PA APIs to detect downtimes.'),
        ('Cost Engine', 'CE-303', 'Calculates MDR (Merchant Discount Rate) for specific card BINs.'),
        ('Decision Engine', 'DE-404', 'The core ML model that selects the routing path.'),
        ('Audit Logger', 'AL-505', 'Asynchronous logger for compliance and debugging.'),
    ]
    
    for name, cid, func in components:
        row_cells = table.add_row().cells
        row_cells[0].text = name
        row_cells[1].text = cid
        row_cells[2].text = func
        
    # Section 3: Performance Benchmarks
    doc.add_heading('3. Performance Benchmarks', level=1)
    doc.add_paragraph("The following targets must be maintained for SLA compliance:")
    
    perf_data = [
        "Maximum P99 Routing Latency: 45ms",
        "Throughput Capacity: 10,000 Transactions Per Second (TPS)",
        "Minimum Decision Accuracy: 99.92%",
        "Failover Recovery Time: < 2 seconds"
    ]
    for item in perf_data:
        doc.add_paragraph(item, style='List Bullet')
        
    # Section 4: Error Codes
    doc.add_heading('4. Standardized Error Codes', level=1)
    err_table = doc.add_table(rows=1, cols=2)
    err_table.style = 'Table Grid'
    err_hdrs = err_table.rows[0].cells
    err_hdrs[0].text = 'Code'
    err_hdrs[1].text = 'Description'
    
    errors = [
        ('ERR_001', 'Provider Downtime Detected'),
        ('ERR_002', 'BIN Range Not Supported'),
        ('ERR_003', 'Currency Conversion Failure'),
        ('ERR_004', 'Invalid Checksum/MAC Signature'),
        ('ERR_005', 'Merchant Velocity Limit Exceeded'),
    ]
    for code, desc in errors:
        row_cells = err_table.add_row().cells
        row_cells[0].text = code
        row_cells[1].text = desc

    # Section 5: Security Requirements
    doc.add_heading('5. Security & Compliance', level=1)
    doc.add_paragraph(
        "All routing decisions containing PII or card metadata must be encrypted "
        "using AES-256-GCM. Keys are rotated every 90 days via HashiCorp Vault. "
        "The system must maintain PCI-DSS Level 1 compliance at all times."
    )
    
    doc.save(path)

if __name__ == "__main__":
    create_complex_docx('/home/carl/Desktop/wisipay/rag-pdf/sample_docs/wisipay_tech_spec.docx')
    print("DOCX created successfully.")
